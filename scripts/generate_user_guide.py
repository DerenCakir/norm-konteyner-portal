"""
Norm Konteyner Sayım Portalı — Kullanıcı Kılavuzu DOCX üretici.

Çalıştır:
    python scripts/generate_user_guide.py

Çıktı: docs/Norm_Konteyner_Kullanici_Kilavuzu.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


PRIMARY = RGBColor(0x1F, 0x3A, 0x8A)
ACCENT = RGBColor(0x25, 0x63, 0xEB)
SUCCESS = RGBColor(0x16, 0xA3, 0x4A)
SUCCESS_BG = "D5F0DD"
DANGER = RGBColor(0xDC, 0x26, 0x26)
DANGER_BG = "FCE4E4"
WARNING_BG = "FFF3CD"
WARNING = RGBColor(0xCA, 0x8A, 0x04)
INFO_BG = "E0EAF8"
MUTED = RGBColor(0x47, 0x55, 0x69)
TEXT = RGBColor(0x0F, 0x17, 0x2A)
SUBTLE = RGBColor(0x94, 0xA3, 0xB8)


def shade_cell(cell, hex_fill: str) -> None:
    """Hücreye arka plan rengi ata (python-docx native desteği yok)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def set_cell_borders(cell, color: str = "DDDDDD", size: int = 4) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


def setup_heading_styles(doc: Document) -> None:
    """Heading 1/2/3 stillerini özelleştir — kurumsal mavi tonlar."""
    styles = doc.styles
    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(20)
    h1.font.bold = True
    h1.font.color.rgb = PRIMARY
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = ACCENT
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = TEXT
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(2)


def add_para(doc, text: str, *, size: float = 11, color: RGBColor = TEXT,
             italic: bool = False, bold: bool = False,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_after: float = 6):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.italic = italic
    run.bold = bold
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.4
    return p


def add_bullet(doc, text: str, *, size: float = 11):
    p = doc.add_paragraph(style="List Bullet")
    if p.runs:
        for r in p.runs:
            r.text = ""
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = TEXT
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.3
    return p


def add_numbered(doc, text: str, *, size: float = 11):
    p = doc.add_paragraph(style="List Number")
    if p.runs:
        for r in p.runs:
            r.text = ""
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = TEXT
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.3
    return p


def add_callout(doc, label: str, body: str, kind: str = "info"):
    """Renkli kutu — info/success/warning/danger.

    Tek hücreli bir tablo olarak çiziyoruz; arka plan + sol kenarlık.
    """
    fill_map = {
        "info": INFO_BG, "success": SUCCESS_BG,
        "warning": WARNING_BG, "danger": DANGER_BG,
    }
    accent_map = {
        "info": ACCENT, "success": SUCCESS,
        "warning": WARNING, "danger": DANGER,
    }
    fill = fill_map.get(kind, INFO_BG)
    accent = accent_map.get(kind, ACCENT)

    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_borders(cell, color="FFFFFF", size=2)
    # Sol kenarda renkli çubuk için ayrı border
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    left = borders.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        borders.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:color"), f"{accent.rgb:06X}" if hasattr(accent, "rgb") else "2563EB")

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3
    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.size = Pt(10.5)
    label_run.font.color.rgb = accent
    body_run = p.add_run(body)
    body_run.font.size = Pt(10.5)
    body_run.font.color.rgb = TEXT

    # Kutuda extra alt boşluk
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def add_section_break(doc):
    """Bölümler arasında ince yatay çizgi efekti — tek hücreli tablo, alt kenar."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "CBD5E1")
    borders.append(bottom)
    tc_pr.append(borders)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)


def add_count_types_table(doc):
    """4 sayım tipini açıklayan tablo."""
    table = doc.add_table(rows=5, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(4.5)
    table.columns[1].width = Cm(11.5)

    # Header
    hdr = table.rows[0]
    for cell, txt in zip(hdr.cells, ["Sayım Tipi", "Açıklama"]):
        shade_cell(cell, "1F3A8A")
        cell.width = Cm(4.5) if txt == "Sayım Tipi" else Cm(11.5)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(txt)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    rows_data = [
        ("Boş", "O renkten kaç adet boş konteyner var?"),
        ("Dolu (toplam)", "O renkten kaç adet dolu konteyner var?"),
        ("Kanban",
         "Dolu konteynerlerin kaç tanesi kanban olarak kullanılıyor? "
         "Kanban dolu'nun ALT KÜMESİDİR — dolu sayısından büyük olamaz."),
        ("Hurdaya Ayrılacak",
         "Artık kullanılmayacak konteynerler (ayağı kırık, çatlak vb.). "
         "Boş ve dolu sayılarına dahil değildir, ayrıca sayılır."),
    ]
    for idx, (label, desc) in enumerate(rows_data, start=1):
        row = table.rows[idx]
        zebra = "F8FAFC" if idx % 2 == 0 else "FFFFFF"
        for cell in row.cells:
            shade_cell(cell, zebra)
            set_cell_borders(cell, color="E2E8F0", size=4)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(11.5)

        p1 = row.cells[0].paragraphs[0]
        p1.paragraph_format.space_before = Pt(4)
        p1.paragraph_format.space_after = Pt(4)
        r1 = p1.add_run(label)
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = PRIMARY

        p2 = row.cells[1].paragraphs[0]
        p2.paragraph_format.space_before = Pt(4)
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(desc)
        r2.font.size = Pt(11)
        r2.font.color.rgb = TEXT
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def add_example_table(doc):
    """Örnek sayım: Mavi renk için 4 değer."""
    table = doc.add_table(rows=2, cols=4)
    table.autofit = False
    for col in table.columns:
        col.width = Cm(4)

    headers = ["Boş", "Dolu (toplam)", "Kanban", "Hurdaya Ayrılacak"]
    for cell, txt in zip(table.rows[0].cells, headers):
        shade_cell(cell, "EEF4FB")
        set_cell_borders(cell, color="C6D4E5", size=4)
        cell.width = Cm(4)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(txt)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = MUTED

    values = ["100", "500", "100", "5"]
    for cell, txt in zip(table.rows[1].cells, values):
        shade_cell(cell, "FFFFFF")
        set_cell_borders(cell, color="E2E8F0", size=4)
        cell.width = Cm(4)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(txt)
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = PRIMARY
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_cover_page(doc):
    # Üstte boşluk
    for _ in range(4):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NORM FASTENERS")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED
    rpr = run._element.get_or_add_rPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), "100")
    rpr.append(spacing)
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Konteyner Sayım Portalı")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = PRIMARY
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Kullanıcı Kılavuzu")
    run.font.size = Pt(16)
    run.font.color.rgb = ACCENT
    run.italic = True
    p.paragraph_format.space_after = Pt(36)

    # Açıklama bloğu — ortalı kutu
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.width = Cm(13)
    shade_cell(cell, "F8FAFC")
    set_cell_borders(cell, color="CBD5E1", size=4)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.4
    run = p.add_run(
        "Bu kılavuz, haftalık konteyner sayım girişlerinizi portal üzerinden "
        "nasıl yapacağınızı adım adım anlatır. Sisteme erişim, form doldurma "
        "kuralları ve sık karşılaşılan sorunlar dahildir."
    )
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT

    # Page break
    doc.add_page_break()


def add_toc(doc):
    """İçindekiler tablosu — Word açtığında F9 ile güncellemesi gerekir."""
    h = doc.add_paragraph()
    h.style = doc.styles["Heading 1"]
    h.text = "İçindekiler"

    p = doc.add_paragraph()
    run = p.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-2" \\h \\z \\u'
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "İçindekiler tablosunu güncellemek için F9 tuşuna basın."
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._element.append(fld_char1)
    run._element.append(instr_text)
    run._element.append(fld_char2)
    run._element.append(placeholder)
    run._element.append(fld_char3)

    doc.add_page_break()


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Norm Fasteners — Konteyner Operasyon Merkezi")
    run.font.size = Pt(9)
    run.font.color.rgb = SUBTLE
    run.italic = True


# ---------------------------------------------------------------------------
def build_document() -> Document:
    doc = Document()

    # Sayfa düzeni
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.4

    setup_heading_styles(doc)
    add_footer(doc)

    add_cover_page(doc)
    add_toc(doc)

    # =========================================================
    # 1. Giriş ve Şifre
    # =========================================================
    doc.add_paragraph("1. Giriş ve Şifre", style="Heading 1")
    add_para(
        doc,
        "Portala bilgisayardaki tarayıcıdan (Chrome, Edge vb.) erişiyorsunuz. "
        "Yönetici tarafından size iletilen kullanıcı adı ve şifre ile sisteme girersiniz.",
    )

    doc.add_paragraph("Adım adım giriş", style="Heading 2")
    add_numbered(doc, "Tarayıcıyı açın, yöneticinin paylaştığı portal adresini girin.")
    add_numbered(doc, "Açılan ekranda kullanıcı adınızı ve şifrenizi yazın.")
    add_numbered(doc, "“Giriş Yap” butonuna basın.")

    add_callout(
        doc,
        "Yanlış şifre",
        "Sağ üst köşede uyarı görüntülenir. Bilgileri kontrol edip tekrar deneyin.",
        kind="danger",
    )
    add_callout(
        doc,
        "Şifrenizi unuttunuz mu",
        "Yöneticinizle iletişime geçin; şifrenizi sıfırlatabilir.",
        kind="info",
    )

    add_section_break(doc)

    # =========================================================
    # 2. Ana Sayfa
    # =========================================================
    doc.add_paragraph("2. Ana Sayfa", style="Heading 1")
    add_para(
        doc,
        "Login sonrası ilk açılan sayfa. Üstte “Hoş geldin, [adınız]” yazısı, "
        "altında bu haftanın tarihi ve sayım giriş penceresinin durumu yer alır.",
    )

    doc.add_paragraph("Pencere durumu", style="Heading 2")
    add_bullet(doc, "Yeşil “Sayım Açık”: o anda sayım giriyor olabilirsiniz.")
    add_bullet(
        doc,
        "Gri “Sayım Kapalı”: pencere açık değil. Açılma zamanı yine ekranda yazılı "
        "(örn. “Pazartesi 09:00 – 12:00 arasında açılır”).",
    )
    add_callout(
        doc,
        "Sayım penceresi",
        "Her hafta Pazartesi 09:00 – 12:00 arasında açıktır. Yönetici farklı bir "
        "saat ayarlamışsa, sistem ona göre çalışır.",
        kind="info",
    )

    doc.add_paragraph("Hızlı erişim kartları", style="Heading 2")
    add_para(doc, "Ana sayfada doğrudan tıklanabilen iki kısayol vardır:")
    add_bullet(doc, "Sayım Girişi — haftalık sayım buradan yapılır.")
    add_bullet(doc, "Haftalık Durum — bu hafta hangi bölümler girdi, hangileri eksik?")

    add_section_break(doc)

    # =========================================================
    # 3. Sayım Girişi
    # =========================================================
    doc.add_paragraph("3. Sayım Girişi", style="Heading 1")
    add_para(
        doc,
        "Sol menüden “Sayım Girişi” seçeneğine tıklayın. Form yüklenir. Pencere "
        "kapalıyken form pasif olur, yalnızca pencere açıkken doldurulabilir.",
    )

    doc.add_paragraph("Hafta ve bölüm seçimi", style="Heading 2")
    add_bullet(doc, "Hafta otomatik gelir, değiştirilemez.")
    add_bullet(
        doc,
        "“Bölüm” açılır listesinden kendi sorumlu olduğunuz bölümü seçin. Birden "
        "fazla bölümden sorumluysanız her biri için ayrı ayrı sayım girersiniz.",
    )

    doc.add_paragraph("Tarih ve saat", style="Heading 2")
    add_para(doc, "Otomatik damgalanır, manuel girilmez. Sistem kayıt anının zamanını yazar.")

    doc.add_paragraph("Yarı mamül tonajı (toplam)", style="Heading 2")
    add_callout(
        doc,
        "Zorunlu alan",
        "Bu alan boş bırakılırsa kayıt yapılamaz. Sistem uyarı verir, kutu kırmızı yanar.",
        kind="danger",
    )
    add_para(
        doc,
        "Bölümünüzde o haftaki yarı mamül tonajını yazın. Ondalık değer girilebilir "
        "(örn. 1234 ya da 12.5).",
    )

    doc.add_paragraph("Renk × Sayım tablosu", style="Heading 2")
    add_para(doc, "Her renk için 4 değer girilir:")
    add_count_types_table(doc)

    add_para(doc, "Örnek — Mavi renk için doğru girilmiş bir sayım:", bold=True, size=11)
    add_example_table(doc)

    add_callout(
        doc,
        "Yorum",
        "100 boş, 500 dolu, doluların 100’ü kanban (doluya dahil), 5 hurdaya "
        "ayrılacak (boş/doluya dahil değil, ayrı sayılır).",
        kind="success",
    )
    add_callout(
        doc,
        "Önemli kural",
        "Kanban miktarı dolu miktarından büyük olamaz. Sistem buna izin vermez.",
        kind="danger",
    )

    doc.add_paragraph("Kaydetme", style="Heading 2")
    add_numbered(doc, "Tüm değerleri girdikten sonra alttaki “Kaydet” butonuna basın.")
    add_numbered(doc, "Sağ üst köşede yeşil “Sayım kaydedildi” bildirimi görünür.")

    doc.add_paragraph("Yanlış girdiyseniz — Güncelleme", style="Heading 2")
    add_para(doc, "Aynı bölüm için aynı hafta içinde formu tekrar açabilirsiniz:")
    add_bullet(doc, "Mevcut değerleriniz formda dolu gelir.")
    add_bullet(doc, "Düzeltmek istediğiniz alanı değiştirin.")
    add_bullet(doc, "Buton bu kez “Güncelle” yazar. Basın.")
    add_bullet(doc, "Sağ üstte “Sayım güncellendi” bildirimi görünür.")
    add_callout(
        doc,
        "Hiç değişmediyse",
        "“Değişiklik yapılmadı” uyarısı gelir. Verileriniz olduğu gibi kalır, kayıp olmaz.",
        kind="info",
    )

    add_section_break(doc)

    # =========================================================
    # 4. Haftalık Durum
    # =========================================================
    doc.add_paragraph("4. Haftalık Durum", style="Heading 1")
    add_para(
        doc,
        "Sol menüden “Haftalık Durum” sayfasına gidin. Bu sayfa, seçili haftanın "
        "tüm bölüm bilgilerini özetler.",
    )

    doc.add_paragraph("Üst özet", style="Heading 2")
    add_bullet(doc, "Toplam bölüm sayısı içinde kaç tanesi sayım girdi.")
    add_bullet(doc, "Eksik kalan bölümlerin listesi (sorumlu kullanıcı adıyla).")

    doc.add_paragraph("Bölüm × Renk Matrisi", style="Heading 2")
    add_para(
        doc,
        "Her bölüm için her rengin sayım değerlerini tek bakışta gösterir. "
        "Hücre formatı: Boş / Dolu / Kanban / Hurdaya Ayrılacak.",
    )
    add_callout(
        doc,
        "Örnek hücre",
        "100/500/100/5 → 100 boş, 500 dolu, 100 kanban, 5 hurdaya ayrılacak.",
        kind="info",
    )
    add_para(doc, "En sağda Tonaj sütunu — bölümün o haftaki toplam tonajı.")

    add_section_break(doc)

    # =========================================================
    # 5. Yetkililer
    # =========================================================
    doc.add_paragraph("5. Yetkililer", style="Heading 1")
    add_para(
        doc,
        "Sol menüden “Yetkililer” sayfasına gidin. Hangi bölüme kimin yetkili "
        "olduğunu, her birinin sayım girip girmediğini gösterir. Kontrol amaçlıdır.",
    )

    add_section_break(doc)

    # =========================================================
    # 6. Çıkış
    # =========================================================
    doc.add_paragraph("6. Çıkış Yapma", style="Heading 1")
    add_para(doc, "Sol menünün altındaki “Çıkış Yap” butonuna basın. Login ekranına dönülür.")
    add_callout(
        doc,
        "Güvenlik",
        "Bilgisayardan kalkıyorsanız mutlaka çıkış yapın. Aksi halde başka biri sizin "
        "oturumunuzla işlem yapabilir.",
        kind="warning",
    )

    add_section_break(doc)

    # =========================================================
    # 7. SSS
    # =========================================================
    doc.add_paragraph("7. Sık Karşılaşılan Sorunlar", style="Heading 1")

    doc.add_paragraph("“Sayım Girişi yapamıyorum, form pasif.”", style="Heading 3")
    add_bullet(doc, "Pencere kapalı — Pazartesi 09:00 – 12:00 dışındasınız. Pencere açıldığında tekrar deneyin.")
    add_bullet(doc, "Bölüm yetkisi atanmamış — “Yetkili bölüm bulunamadı” uyarısı görüyorsanız yöneticinizle iletişime geçin.")

    doc.add_paragraph("“Şifremi unuttum.”", style="Heading 3")
    add_para(doc, "Yöneticinize haber verin. Şifrenizi sıfırlayıp size yenisini iletecektir.")

    doc.add_paragraph("“Yanlış değer girdim, düzeltebilir miyim?”", style="Heading 3")
    add_para(
        doc,
        "Pencere hâlâ açıksa: aynı bölüm + aynı hafta için Sayım Girişi sayfasını "
        "tekrar açıp değerleri düzeltin, “Güncelle”ye basın. Pencere kapandıysa "
        "yöneticinize bildirin; gerektiğinde sizin için geç giriş açabilir.",
    )

    doc.add_paragraph("“Saat geçti, sayımı giremedim.”", style="Heading 3")
    add_para(
        doc,
        "Yöneticinize haber verin. Yönetici geç giriş penceresi açabilir. "
        "Açıldığında ekranınızda “Sayım Açık” durumu yine görünür ve normal "
        "şekilde sayım girersiniz.",
    )

    doc.add_paragraph("“Kanban miktarını dolu’dan büyük yazınca uyarı veriyor.”", style="Heading 3")
    add_para(
        doc,
        "Doğru çalışıyor. Kanban konteynerler dolu konteynerlerin alt kümesidir; "
        "doludan büyük olamaz. Önce dolu sayısını gözden geçirin.",
    )

    doc.add_paragraph("“Tonaj alanı kırmızı yanıyor.”", style="Heading 3")
    add_para(
        doc,
        "Tonaj zorunlu alandır, boş bırakılamaz. Değer girip “Kaydet”e tekrar basın.",
    )

    return doc


def main() -> int:
    doc = build_document()
    out_dir = Path(__file__).resolve().parent.parent / "docs"
    out_dir.mkdir(exist_ok=True)
    out_path = out_dir / "Norm_Konteyner_Kullanici_Kilavuzu.docx"
    doc.save(out_path)
    print(f"OK -> {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
